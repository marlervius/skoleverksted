from __future__ import annotations

import html
import io
import logging
import os
import re
from pathlib import Path
from typing import Iterable
from urllib.parse import parse_qsl, urlencode, urlsplit, urlunsplit

from .models import Compendium, CompendiumSource


logger = logging.getLogger(__name__)

DOCUMENT_TYPE_LABELS = {
    "thematic": "Tematisk fordypning",
    "chronological": "Kronologisk oversikt",
    "reference": "Oppslagsverk",
    "comparative": "Sammenlignende kompendium",
    "source_collection": "Kildesamling",
    "appendix": "Appendiks",
}
_TRANSIENT_SOURCE_HOSTS = {"vertexaisearch.cloud.google.com"}
_TRACKING_QUERY_KEYS = {"gclid", "fbclid", "mc_cid", "mc_eid"}
_COMMON_LANGUAGE_FIXES = {
    "produjonsmidlene": "produksjonsmidlene",
    "produjonsmiddel": "produksjonsmiddel",
}
_UNRESOLVED_SCOPE_MARKERS = (
    "ikke spesifisert",
    "avgrenses av læreren",
    "før produksjon",
)


def safe_filename(value: str, suffix: str) -> str:
    stem = re.sub(r"[^a-zA-Z0-9æøåÆØÅ_-]+", "-", value).strip("-").lower()
    return f"{stem[:100] or 'kompendium'}.{suffix}"


def _typst_escape(value: str) -> str:
    value = value.replace("\\", "\\\\")
    value = re.sub(r"([#\[\]$<>@*_~`])", r"\\\1", value)
    return value


def _typst_string(value: str) -> str:
    return value.replace("\\", "\\\\").replace('"', '\\"')


def _clean_markdown(value: str) -> str:
    value = html.unescape(value)
    value = re.sub(r"<br\s*/?>", "; ", value, flags=re.IGNORECASE)
    value = re.sub(r"</?[a-z][^>]*>", "", value, flags=re.IGNORECASE)
    value = re.sub(r"\*\*(.*?)\*\*", r"\1", value)
    value = re.sub(r"__(.*?)__", r"\1", value)
    value = re.sub(r"(?<!\*)\*([^*\n]+)\*(?!\*)", r"\1", value)
    value = re.sub(r"(?<!_)_([^_\n]+)_(?!_)", r"\1", value)
    value = re.sub(r"\[(.*?)\]\((https?://[^)]+)\)", r"\1 (\2)", value)
    for typo, correction in _COMMON_LANGUAGE_FIXES.items():
        value = re.sub(rf"\b{re.escape(typo)}\b", correction, value, flags=re.IGNORECASE)
    value = re.sub(r"\s+([,.;:!?])", r"\1", value)
    value = re.sub(r"[ \t]{2,}", " ", value)
    return value.strip()


def _normalize_markdown_newlines(value: str) -> str:
    """Turn model-produced escaped line breaks into real Markdown lines."""
    return (
        value.replace("\\r\\n", "\n")
        .replace("\\n", "\n")
        .replace("\\r", "\n")
    )


def _table_to_typst(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    columns = max(len(row) for row in rows)
    cells: list[str] = []
    for row_index, row in enumerate(rows):
        padded = row + [""] * (columns - len(row))
        for value in padded:
            content = _typst_escape(_clean_markdown(value))
            if row_index == 0:
                cells.append(f"[#strong[{content}]]")
            else:
                cells.append(f"[{content}]")
    return (
        f"#table(columns: {columns}, inset: 5pt, stroke: 0.35pt + rgb(\"#D8DDE5\"),\n"
        + ",\n".join(cells)
        + ")\n#v(8pt)"
    )


def markdown_to_typst(markdown: str, chapter_title: str = "") -> str:
    lines = _normalize_markdown_newlines(markdown).replace("\r\n", "\n").splitlines()
    output: list[str] = []
    index = 0
    first_heading_skipped = False
    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines: list[str] = []
            while index < len(lines):
                candidate = lines[index].strip()
                if not (candidate.startswith("|") and candidate.endswith("|")):
                    break
                table_lines.append(candidate)
                index += 1
            rows = [
                [cell.strip() for cell in row.strip("|").split("|")]
                for row in table_lines
                if not re.fullmatch(r"\|?[\s:|-]+\|?", row)
            ]
            output.append(_table_to_typst(rows))
            continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            title = _clean_markdown(heading.group(2))
            if not first_heading_skipped and chapter_title and title.casefold() == chapter_title.casefold():
                first_heading_skipped = True
            else:
                source_level = len(heading.group(1)) - (1 if chapter_title else 0)
                level = min(3, max(2, source_level))
                output.append(f"{'=' * level} {_typst_escape(title)}")
            index += 1
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet:
            output.append(f"- {_typst_escape(_clean_markdown(bullet.group(1)))}")
            index += 1
            continue
        numbered = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if numbered:
            output.append(f"+ {_typst_escape(_clean_markdown(numbered.group(1)))}")
            index += 1
            continue
        if stripped.startswith(">"):
            output.append(
                '#block(fill: rgb("#F4F6F8"), inset: 8pt, stroke: (left: 2pt + rgb("#356A8A")))['
                + _typst_escape(_clean_markdown(stripped.lstrip("> ")))
                + "]"
            )
            index += 1
            continue
        if stripped:
            output.append(_typst_escape(_clean_markdown(stripped)))
        else:
            output.append("#v(5pt)")
        index += 1
    return "\n".join(output)


def _canonical_source_url(value: str) -> str:
    url = value.strip()
    if not url.startswith(("https://", "http://")):
        return ""
    try:
        parsed = urlsplit(url)
    except ValueError:
        return ""
    host = (parsed.hostname or "").lower().removeprefix("www.")
    if not host or host in _TRANSIENT_SOURCE_HOSTS:
        return ""
    query = urlencode([
        (key, item)
        for key, item in parse_qsl(parsed.query, keep_blank_values=True)
        if key.lower() not in _TRACKING_QUERY_KEYS
        and not key.lower().startswith("utm_")
    ])
    path = parsed.path or "/"
    if path != "/":
        path = path.rstrip("/")
    return urlunsplit((parsed.scheme.lower(), parsed.netloc.lower(), path, query, ""))


def _source_identity(source: CompendiumSource) -> str:
    url = _canonical_source_url(source.url)
    if url:
        parsed = urlsplit(url)
        return f"{parsed.hostname or ''}{parsed.path}".casefold().rstrip("/")
    return re.sub(
        r"\W+",
        "",
        f"{source.title}|{source.publisher}".casefold(),
        flags=re.UNICODE,
    )


def _clean_source(source: CompendiumSource) -> CompendiumSource | None:
    title = re.sub(r"\s+", " ", html.unescape(source.title)).strip(" –—-|")
    publisher = re.sub(r"\s+", " ", html.unescape(source.publisher)).strip(" –—-|")
    raw_url = source.url.strip()
    url = _canonical_source_url(raw_url)
    if raw_url and not url:
        return None
    if not title:
        return None
    return CompendiumSource(title=title, publisher=publisher, url=url)


def _unique_sources(compendium: Compendium) -> list[CompendiumSource]:
    result: list[CompendiumSource] = []
    seen: set[str] = set()
    for chapter in compendium.chapters:
        for source in chapter.sources:
            cleaned = _clean_source(source)
            if cleaned is None:
                continue
            key = _source_identity(cleaned)
            if key and key not in seen:
                seen.add(key)
                result.append(cleaned)
    return result


def _is_meaningful_scope_value(value: str) -> bool:
    normalized = re.sub(r"\s+", " ", value).strip().casefold()
    return bool(normalized) and not any(marker in normalized for marker in _UNRESOLVED_SCOPE_MARKERS)


def _student_facing_scope_note(value: str) -> str:
    sentences = re.split(r"(?<=[.!?])\s+", re.sub(r"\s+", " ", value).strip())
    kept = [
        sentence
        for sentence in sentences
        if sentence
        and "læreren" not in sentence.casefold()
        and "før bruk" not in sentence.casefold()
        and "før produksjon" not in sentence.casefold()
    ]
    return " ".join(kept)


def _student_facing_scope_items(values: Iterable[str]) -> list[str]:
    result: list[str] = []
    for value in values:
        normalized = value.casefold()
        if any(marker in normalized for marker in ("ta bare med", "presenteres ikke", "læreren", "før produksjon")):
            continue
        cleaned = _clean_markdown(value)
        if cleaned and cleaned not in result:
            result.append(cleaned)
    return result


def _scope_label(value: str) -> str:
    return {
        "complete": "Fullstendig innenfor kriteriene",
        "documented": "Dokumentert oversikt",
        "selected": "Faglig utvalg",
    }.get(value, "Faglig utvalg")


def build_typst_document(compendium: Compendium, *, image_path: str = "", image_credit: str = "") -> str:
    scope = compendium.scope_contract
    chapters = sorted(compendium.chapters, key=lambda item: item.order)
    source_items = _unique_sources(compendium)
    document_type = DOCUMENT_TYPE_LABELS.get(compendium.kind, "Faglig kompendium")
    image_block = ""
    if image_path and Path(image_path).is_file():
        image_name = _typst_escape(Path(image_path).name)
        image_block = f"""
#v(10mm)
#figure(
  image("{image_name}", width: 78%),
  caption: [{_typst_escape(image_credit or "Pedagogisk illustrasjon")}],
)
"""
    chapter_blocks: list[str] = []
    for chapter in chapters:
        content = markdown_to_typst(chapter.content_markdown, chapter.title)
        guiding_questions = ""
        if chapter.guiding_questions:
            question_items = "\n".join(
                f"+ {_typst_escape(_clean_markdown(question))}"
                for question in chapter.guiding_questions[:4]
            )
            guiding_questions = f"""
#block(fill: rgb("#EDF7F7"), inset: 10pt, radius: 5pt, stroke: 0.5pt + rgb("#B8D7D8"))[
  #strong[Spørsmål å ha med seg]
  #v(3pt)
  {question_items}
]
#v(8pt)
"""
        key_facts = ""
        if chapter.key_facts:
            fact_items = "\n".join(
                f"- {_typst_escape(_clean_markdown(item))}"
                for item in chapter.key_facts[:6]
            )
            key_facts = f"""
#block(breakable: false, fill: rgb("#F4F6F8"), inset: 9pt, radius: 5pt)[
  #set text(size: 9.2pt)
  #set par(justify: false, leading: 0.62em)
  #text(size: 12pt, weight: 600, fill: rgb("#235F72"))[Kort oppsummert]
  #v(4pt)
  {fact_items}
]
"""
        glossary = ""
        if compendium.include_glossary and chapter.glossary:
            visible_glossary = chapter.glossary[:12]
            terms = "\n".join(
                f"- {_typst_escape(_clean_markdown(item))}"
                for item in visible_glossary
            )
            breakable = "false" if len(visible_glossary) <= 8 else "true"
            glossary = f"""
#block(breakable: {breakable}, inset: (top: 4pt, bottom: 2pt))[
  #set text(size: 9.2pt)
  #set par(justify: false, leading: 0.62em)
  #text(size: 12pt, weight: 600, fill: rgb("#235F72"))[Begreper]
  #v(4pt)
  {terms}
]
"""
        chapter_blocks.append(
            f"= {_typst_escape(chapter.title)}\n"
            f"#text(fill: rgb(\"#5A6572\"), style: \"italic\")[{_typst_escape(chapter.purpose)}]\n"
            f"#v(8pt)\n{guiding_questions}\n{content}\n{key_facts}\n{glossary}"
        )

    scope_rows: list[str] = []
    if _is_meaningful_scope_value(scope.reference_date):
        scope_rows.extend(("[#strong[Tidsrom]],", f"[{_typst_escape(scope.reference_date)}],"))
    if _is_meaningful_scope_value(scope.geography):
        scope_rows.extend(("[#strong[Geografisk ramme]],", f"[{_typst_escape(scope.geography)}],"))
    scope_rows.extend(("[#strong[Dekning]],", f"[{_typst_escape(_scope_label(scope.completeness_label))}],"))
    scope_table = (
        "#table(\n"
        "  columns: (38mm, 1fr),\n"
        "  inset: 6pt,\n"
        "  stroke: 0.35pt + rgb(\"#D8DDE5\"),\n"
        + "\n".join(scope_rows)
        + "\n)"
    )
    scope_note = _student_facing_scope_note(scope.completeness_note)
    inclusion_items = _student_facing_scope_items(scope.inclusion_criteria)
    exclusion_items = _student_facing_scope_items(scope.exclusions)
    scope_details = ""
    if inclusion_items or exclusion_items:
        sections: list[str] = []
        if inclusion_items:
            sections.append(
                "=== Dette behandles\n"
                + "\n".join(f"- {_typst_escape(item)}" for item in inclusion_items)
            )
        if exclusion_items:
            sections.append(
                "=== Avgrensninger\n"
                + "\n".join(f"- {_typst_escape(item)}" for item in exclusion_items)
            )
        scope_details = "\n\n".join(sections)

    bibliography_lines: list[str] = []
    for source in source_items:
        label = _typst_escape(source.title)
        if source.publisher and source.publisher.casefold() not in source.title.casefold():
            label += f" — {_typst_escape(source.publisher)}"
        if source.url:
            bibliography_lines.append(
                f'+ #link("{_typst_string(source.url)}")[{label}]'
            )
        else:
            bibliography_lines.append(f"+ {label}")
    bibliography = "\n".join(bibliography_lines) or (
        "Det er ikke registrert eksterne kilder i denne versjonen."
    )
    truth_lines: list[str] = []
    for chapter in chapters:
        passport = chapter.truth_passport
        if not passport:
            continue
        removed = len(passport.removed_claims)
        handling = (
            f" {removed} udokumentert(e) påstand(er) ble fjernet før ferdigstilling."
            if removed
            else ""
        )
        truth_lines.append(
            f"+ #strong[{_typst_escape(chapter.title)}]: "
            f"{passport.verified_claims} av {passport.total_claims} registrerte "
            f"faktapåstander ble dokumentert med konkrete kilder.{handling}"
        )
    truth_summary = "\n".join(truth_lines) or (
        "Det følger ikke et maskinlesbart faktapass med denne versjonen."
    )

    learning_goals = ""
    if compendium.competency_goals:
        goals = "\n".join(
            f"- {_typst_escape(_clean_markdown(goal))}"
            for goal in compendium.competency_goals
        )
        learning_goals = f"""
== Mål for arbeidet
{goals}
"""

    chapter_preview = " · ".join(chapter.title for chapter in chapters[:5])
    if len(chapters) > 5:
        chapter_preview += " · …"
    tasks = ""
    if compendium.include_reflection_tasks:
        tasks = f"""
= Videre arbeid

#block(fill: rgb("#F3F0FA"), inset: 12pt, radius: 5pt, stroke: 0.5pt + rgb("#D5C9EB"))[
+ Hvilke deler av framstillingen er best dokumentert, og hvor er usikkerheten størst?
+ Sammenlign to aktører, områder eller perioder fra kompendiet.
+ Velg én kilde fra litteraturlisten og vurder hva den kan og ikke kan fortelle.
+ Formuler et nytt fordypningsspørsmål som kompendiet ikke besvarer.
]
"""

    return f"""
#set document(title: "{_typst_escape(compendium.title)}", author: "Skoleverksted")
#set page(
  paper: "a4",
  margin: (top: 22mm, bottom: 20mm, left: 22mm, right: 18mm),
  header: context if counter(page).get().first() > 2 [
    #set text(size: 8pt, fill: rgb("#687583"))
    {_typst_escape(compendium.subject)} · {_typst_escape(document_type)}
    #h(1fr) {_typst_escape(compendium.level)}
    #line(length: 100%, stroke: 0.35pt + rgb("#D8DDE5"))
  ],
  footer: context if counter(page).get().first() > 1 [
    #line(length: 100%, stroke: 0.35pt + rgb("#D8DDE5"))
    #set text(size: 8pt, fill: rgb("#87919B"))
    Skoleverksted #h(1fr) Side #counter(page).display()
  ],
)
#set text(font: ("Source Sans 3", "Noto Sans", "Liberation Sans"), size: 10.5pt, lang: "nb", fill: rgb("#22272B"))
#set par(justify: true, leading: 0.68em)
#set heading(numbering: "1.1")
#show heading.where(level: 1): it => block(above: 20pt, below: 10pt)[
  #text(size: 19pt, weight: 600, fill: rgb("#174D6C"))[#it.body]
  #line(length: 100%, stroke: 1.2pt + rgb("#3E8E9B"))
]
#show heading.where(level: 2): it => block(above: 13pt, below: 6pt)[
  #text(size: 13pt, weight: 600, fill: rgb("#235F72"))[#it.body]
]
#show heading.where(level: 3): it => block(above: 10pt, below: 4pt)[
  #text(size: 11pt, weight: 600)[#it.body]
]

#align(center)[
  #v(10mm)
  #block(width: 100%, fill: rgb("#173F55"), inset: (x: 16mm, y: 15mm), radius: 8pt)[
    #set text(fill: rgb("#FFFFFF"))
    #set par(justify: false)
    #text(size: 9pt, weight: 600, fill: rgb("#8ED4D8"))[SKOLEVERKSTED · {_typst_escape(compendium.subject.upper())}]
    #v(7mm)
    #text(size: 29pt, weight: 650, hyphenate: false)[{_typst_escape(compendium.title)}]
    #v(5mm)
    #text(size: 12pt, fill: rgb("#D7E6EA"))[{_typst_escape(compendium.level)} · {_typst_escape(compendium.audience)}]
  ]
  #v(8mm)
  #block(width: 86%, fill: rgb("#EFF5F6"), inset: 13pt, radius: 6pt)[
    #set par(justify: false)
    #text(size: 9pt, weight: 600, fill: rgb("#3E8E9B"))[{_typst_escape(document_type.upper())}]
    #v(3pt)
    #text(size: 10.5pt)[{_typst_escape(compendium.purpose or compendium.topic)}]
    #v(6pt)
    #text(size: 9pt, fill: rgb("#596875"))[{_typst_escape(chapter_preview)}]
  ]
  {image_block}
]
#pagebreak()

= Om ressursen

Dette kompendiet gir en {_typst_escape(_scope_label(scope.completeness_label).lower())}
av {_typst_escape(compendium.topic)}. Bruk overskriftene, spørsmålene og
oppsummeringene til å få oversikt før du går inn i detaljene.

== Ramme for framstillingen

{scope_table}

{_typst_escape(scope_note)}

{scope_details}

{learning_goals}

== Slik kan du arbeide

#grid(
  columns: (1fr, 1fr, 1fr),
  gutter: 8pt,
  [#block(fill: rgb("#EDF7F7"), inset: 9pt, radius: 4pt)[#strong[1. Orienter deg]\nLes spørsmålene før hvert kapittel.]],
  [#block(fill: rgb("#F4F6F8"), inset: 9pt, radius: 4pt)[#strong[2. Finn sammenhenger]\nBruk tabeller og begreper aktivt.]],
  [#block(fill: rgb("#F3F0FA"), inset: 9pt, radius: 4pt)[#strong[3. Kontroller]\nUndersøk kildene og vurder påstandene.]],
)

#pagebreak()
#outline(title: [Innhold], depth: 1, indent: auto)
#pagebreak()

{"\n#v(10pt)\n".join(chapter_blocks)}

{tasks}

= Kilder og videre lesning

{bibliography}

== Faktagrunnlag

Faktapasset registrerer etterprøvbare påstander og kilder. Grønn status betyr
at udokumenterte påstander er fjernet eller tydelig nyansert før ferdigstilling;
det er ikke en garanti for at enhver formulering er uangripelig.

{truth_summary}
"""


def _add_markdown_to_docx(document, markdown: str, chapter_title: str) -> None:
    lines = _normalize_markdown_newlines(markdown).replace("\r\n", "\n").splitlines()
    first_heading_skipped = False
    index = 0
    while index < len(lines):
        raw = lines[index].strip()
        if raw.startswith("|") and raw.endswith("|"):
            table_lines: list[str] = []
            while index < len(lines):
                candidate = lines[index].strip()
                if not (candidate.startswith("|") and candidate.endswith("|")):
                    break
                table_lines.append(candidate)
                index += 1
            rows = [
                [_clean_markdown(cell.strip()) for cell in row.strip("|").split("|")]
                for row in table_lines
                if not re.fullmatch(r"\|?[\s:|-]+\|?", row)
            ]
            if rows:
                columns = max(len(row) for row in rows)
                table = document.add_table(rows=len(rows), cols=columns)
                table.style = "Table Grid"
                for row_index, row in enumerate(rows):
                    for column_index, value in enumerate(row):
                        cell = table.cell(row_index, column_index)
                        cell.text = value
                        if row_index == 0:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
            continue
        stripped = _clean_markdown(raw)
        if not stripped:
            index += 1
            continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            title = _clean_markdown(heading.group(2))
            if not first_heading_skipped and title.casefold() == chapter_title.casefold():
                first_heading_skipped = True
                continue
            source_level = len(heading.group(1)) - (1 if chapter_title else 0)
            document.add_heading(title, level=min(3, max(2, source_level)))
        elif re.match(r"^[-*]\s+", stripped):
            document.add_paragraph(re.sub(r"^[-*]\s+", "", stripped), style="List Bullet")
        elif re.match(r"^\d+[.)]\s+", stripped):
            document.add_paragraph(re.sub(r"^\d+[.)]\s+", "", stripped), style="List Number")
        elif re.fullmatch(r"\|?[\s:|-]+\|?", stripped):
            continue
        else:
            document.add_paragraph(stripped.strip("| "))
        index += 1


def _add_docx_hyperlink(paragraph, text: str, url: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE

    relationship_id = paragraph.part.relate_to(
        url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "235F72")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    run.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def build_docx(compendium: Compendium, *, image_path: str = "", image_credit: str = "") -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    title = document.add_heading(compendium.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(f"{compendium.subject} · {compendium.level} · {compendium.audience}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if compendium.purpose:
        purpose = document.add_paragraph(compendium.purpose)
        purpose.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if image_path and Path(image_path).is_file():
        document.add_picture(image_path, width=Inches(5.5))
        caption = document.add_paragraph(image_credit or "Pedagogisk illustrasjon")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()

    document.add_heading("Om ressursen", level=1)
    scope = compendium.scope_contract
    document.add_paragraph(
        f"Dette kompendiet gir en {_scope_label(scope.completeness_label).lower()} "
        f"av {compendium.topic}."
    )
    scope_rows: list[tuple[str, str]] = []
    if _is_meaningful_scope_value(scope.reference_date):
        scope_rows.append(("Tidsrom", scope.reference_date))
    if _is_meaningful_scope_value(scope.geography):
        scope_rows.append(("Geografisk ramme", scope.geography))
    scope_rows.append(("Dekning", _scope_label(scope.completeness_label)))
    table = document.add_table(rows=len(scope_rows), cols=2)
    table.style = "Table Grid"
    for row, values in zip(table.rows, scope_rows):
        row.cells[0].text, row.cells[1].text = values
    scope_note = _student_facing_scope_note(scope.completeness_note)
    if scope_note:
        document.add_paragraph(scope_note)
    if compendium.competency_goals:
        document.add_heading("Mål for arbeidet", level=2)
        for goal in compendium.competency_goals:
            document.add_paragraph(_clean_markdown(goal), style="List Bullet")
    document.add_heading("Slik kan du arbeide", level=2)
    for item in (
        "Les spørsmålene før hvert kapittel.",
        "Bruk tabeller, oppsummeringer og begreper til å finne sammenhenger.",
        "Undersøk kildene og vurder hvilke påstander de støtter.",
    ):
        document.add_paragraph(item, style="List Number")
    document.add_page_break()

    for chapter in sorted(compendium.chapters, key=lambda item: item.order):
        document.add_heading(chapter.title, level=1)
        if chapter.purpose:
            paragraph = document.add_paragraph(chapter.purpose)
            paragraph.style = document.styles["Quote"]
        if chapter.guiding_questions:
            document.add_heading("Spørsmål å ha med seg", level=2)
            for question in chapter.guiding_questions[:4]:
                document.add_paragraph(_clean_markdown(question), style="List Number")
        _add_markdown_to_docx(document, chapter.content_markdown, chapter.title)
        if chapter.key_facts:
            document.add_heading("Kort oppsummert", level=2)
            for item in chapter.key_facts[:6]:
                document.add_paragraph(_clean_markdown(item), style="List Bullet")
        if compendium.include_glossary and chapter.glossary:
            document.add_heading("Begreper", level=2)
            for item in chapter.glossary:
                document.add_paragraph(_clean_markdown(item), style="List Bullet")

    if compendium.include_reflection_tasks:
        document.add_heading("Videre arbeid", level=1)
        for task in (
            "Hvilke deler av framstillingen er best dokumentert, og hvor er usikkerheten størst?",
            "Sammenlign to aktører, områder eller perioder fra kompendiet.",
            "Velg én kilde fra litteraturlisten og vurder hva den kan og ikke kan fortelle.",
            "Formuler et nytt fordypningsspørsmål som avgrensningskontrakten ikke dekker.",
        ):
            document.add_paragraph(task, style="List Number")

    document.add_heading("Kilder og videre lesning", level=1)
    sources = _unique_sources(compendium)
    if sources:
        for source in sources:
            value = source.title
            if source.publisher:
                value += f" — {source.publisher}"
            if source.url:
                paragraph = document.add_paragraph(style="List Number")
                _add_docx_hyperlink(paragraph, value, source.url)
            else:
                document.add_paragraph(value, style="List Number")
    else:
        document.add_paragraph("Det er ikke registrert eksterne kilder i denne versjonen.")

    document.add_heading("Faktagrunnlag", level=2)
    document.add_paragraph(
        "Faktapasset registrerer etterprøvbare påstander og kilder. Grønn status "
        "betyr at udokumenterte påstander er fjernet eller tydelig nyansert før "
        "ferdigstilling; det er ikke en garanti for at enhver formulering er uangripelig."
    )
    passports = [
        (chapter.title, chapter.truth_passport)
        for chapter in sorted(compendium.chapters, key=lambda item: item.order)
        if chapter.truth_passport
    ]
    if passports:
        for chapter_title, passport in passports:
            value = (
                f"{chapter_title}: {passport.verified_claims} av "
                f"{passport.total_claims} registrerte faktapåstander ble "
                "dokumentert med konkrete kilder."
            )
            if passport.removed_claims:
                value += (
                    f" {len(passport.removed_claims)} udokumentert(e) "
                    "påstand(er) ble fjernet før ferdigstilling."
                )
            document.add_paragraph(value, style="List Bullet")
    else:
        document.add_paragraph(
            "Det følger ikke et maskinlesbart faktapass med denne versjonen."
        )

    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def render_compendium(
    compendium: Compendium,
    *,
    image_path: str = "",
    image_credit: str = "",
) -> tuple[bytes, bytes, str, str]:
    from VGS_KI.backend.pdf_service import compile_typst

    try:
        typst = build_typst_document(compendium, image_path=image_path, image_credit=image_credit)
        pdf = compile_typst(typst, image_path=image_path or None)
        docx = build_docx(compendium, image_path=image_path, image_credit=image_credit)
    except Exception as exc:
        if not image_path:
            raise
        logger.warning(
            "Kompendiumbildet kunne ikke bygges inn; lager dokumentene uten bilde: %s",
            exc,
        )
        typst = build_typst_document(compendium)
        pdf = compile_typst(typst)
        docx = build_docx(compendium)
    return (
        pdf,
        docx,
        safe_filename(compendium.title, "pdf"),
        safe_filename(compendium.title, "docx"),
    )
