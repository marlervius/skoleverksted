"""Generate .docx files from lesson content using python-docx."""
import re
from io import BytesIO
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches


def _add_run_with_inline_bold(paragraph, text: str) -> None:
    """Parse **bold** markers and add runs to a paragraph."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            paragraph.add_run(part)


def _parse_and_add_text(doc: Document, text: str) -> None:
    """Convert markdown-ish agent output into Word paragraphs."""
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Blank line → paragraph break (skip)
        if not line.strip():
            i += 1
            continue

        # ### **Heading** or ## **Heading**
        heading_match = re.match(r"^(#{2,4})\s*\*?\*?(.*?)\*?\*?\s*$", line)
        if heading_match:
            level = len(heading_match.group(1))          # 2, 3, or 4
            heading_text = heading_match.group(2).strip()
            word_level = min(level, 4)                   # Word supports H1-H9
            doc.add_heading(heading_text, level=word_level)
            i += 1
            continue

        # Bullet / numbered list item
        if re.match(r"^[-•]\s+", line):
            content = re.sub(r"^[-•]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            _add_run_with_inline_bold(p, content)
            i += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            content = re.sub(r"^\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            _add_run_with_inline_bold(p, content)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_run_with_inline_bold(p, line.strip())
        i += 1


def _add_section_heading(doc: Document, title: str, color: Optional[RGBColor] = None) -> None:
    """Add a visually distinct section separator heading."""
    doc.add_paragraph()
    p = doc.add_heading(title, level=1)
    if color:
        for run in p.runs:
            run.font.color.rgb = color


def create_lesson_docx(
    content_text: str,
    worksheet_text: str,
    topic: str,
    level: str,
    subject: str,
    faktarapport: Optional[str] = None,
) -> bytes:
    """Generate a .docx file from lesson content. Returns raw bytes."""
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── Title ─────────────────────────────────────────────────────────────────
    title_para = doc.add_heading(topic, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"{subject}  ·  {level}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)  # slate-500

    doc.add_paragraph()

    # ── Fagtekst ─────────────────────────────────────────────────────────────
    _add_section_heading(doc, "Fagtekst", color=RGBColor(0x1E, 0x40, 0xAF))  # blue-800
    _parse_and_add_text(doc, content_text)

    # ── Page break before worksheet ───────────────────────────────────────────
    doc.add_page_break()

    # ── Arbeidsark ────────────────────────────────────────────────────────────
    _add_section_heading(doc, "Arbeidsark", color=RGBColor(0x06, 0x6E, 0x4A))  # emerald-800
    if worksheet_text and worksheet_text.strip():
        _parse_and_add_text(doc, worksheet_text)
    else:
        doc.add_paragraph("(Ingen arbeidsark generert)")

    # ── Faktarapport (optional) ───────────────────────────────────────────────
    if faktarapport and faktarapport.strip():
        doc.add_page_break()
        _add_section_heading(doc, "Faktarapport — kun for læreren", color=RGBColor(0x92, 0x40, 0x0E))  # amber-800
        note = doc.add_paragraph()
        run = note.add_run("Dette avsnittet er ment for læreren og skal ikke deles med elevene.")
        run.italic = True
        run.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
        doc.add_paragraph()
        _parse_and_add_text(doc, faktarapport)

    # ── Footer note ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_run = footer_p.add_run("Generert av VGS Lærerassistent")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)  # slate-400
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
