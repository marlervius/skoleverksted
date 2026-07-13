"""
Word Document Exporter for MateMaTeX.
Converts LaTeX content to Microsoft Word (.docx) format.
"""

import re
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def is_word_export_available() -> bool:
    """Check if Word export is available."""
    return DOCX_AVAILABLE


def latex_to_word(latex_content: str, output_path: str) -> Optional[str]:
    """
    Convert LaTeX content to a Word document.
    
    Args:
        latex_content: The LaTeX source code.
        output_path: Path for the output .docx file.
    
    Returns:
        Path to the generated Word document, or None if failed.
    """
    if not DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is not installed. "
            "Run: pip install python-docx"
        )
    
    doc = Document()
    
    # Set up styles
    _setup_styles(doc)
    
    # Extract title
    title_match = re.search(r'\\title\{([^}]+)\}', latex_content)
    if title_match:
        title = _clean_latex(title_match.group(1))
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Extract content between \begin{document} and \end{document}
    doc_match = re.search(
        r'\\begin\{document\}(.*)\\end\{document\}',
        latex_content,
        re.DOTALL
    )
    
    if doc_match:
        content = doc_match.group(1)
    else:
        content = latex_content
    
    # Remove \maketitle and similar commands
    content = re.sub(r'\\maketitle', '', content)
    content = re.sub(r'\\tableofcontents', '', content)
    
    # Process content
    _process_content(doc, content)
    
    # Save document
    output_path = str(output_path)
    if not output_path.endswith('.docx'):
        output_path += '.docx'
    
    doc.save(output_path)
    return output_path


def _setup_styles(doc: 'Document'):
    """Set up custom styles for the document."""
    styles = doc.styles
    
    # Definition box style
    if 'Definition' not in [s.name for s in styles]:
        def_style = styles.add_style('Definition', WD_STYLE_TYPE.PARAGRAPH)
        def_style.font.size = Pt(11)
        def_style.font.color.rgb = RGBColor(0, 102, 204)
        def_style.paragraph_format.left_indent = Inches(0.25)
        def_style.paragraph_format.space_before = Pt(12)
        def_style.paragraph_format.space_after = Pt(12)
    
    # Example box style
    if 'Example' not in [s.name for s in styles]:
        ex_style = styles.add_style('Example', WD_STYLE_TYPE.PARAGRAPH)
        ex_style.font.size = Pt(11)
        ex_style.font.color.rgb = RGBColor(0, 153, 76)
        ex_style.paragraph_format.left_indent = Inches(0.25)
        ex_style.paragraph_format.space_before = Pt(12)
        ex_style.paragraph_format.space_after = Pt(12)
    
    # Task box style
    if 'Task' not in [s.name for s in styles]:
        task_style = styles.add_style('Task', WD_STYLE_TYPE.PARAGRAPH)
        task_style.font.size = Pt(11)
        task_style.paragraph_format.left_indent = Inches(0.25)
        task_style.paragraph_format.space_before = Pt(8)
        task_style.paragraph_format.space_after = Pt(8)


def _process_content(doc: 'Document', content: str):
    """Process LaTeX content and add to Word document."""
    
    # Split into chunks by sections
    chunks = re.split(r'(\\section\*?\{[^}]+\}|\\subsection\*?\{[^}]+\})', content)
    
    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk:
            continue
        
        # Handle sections
        section_match = re.match(r'\\section\*?\{([^}]+)\}', chunk)
        if section_match:
            title = _clean_latex(section_match.group(1))
            doc.add_heading(title, 1)
            continue
        
        # Handle subsections
        subsection_match = re.match(r'\\subsection\*?\{([^}]+)\}', chunk)
        if subsection_match:
            title = _clean_latex(subsection_match.group(1))
            doc.add_heading(title, 2)
            continue
        
        # Process regular content
        _process_chunk(doc, chunk)


def _process_chunk(doc: 'Document', chunk: str):
    """Process a chunk of LaTeX content."""
    
    # Handle definitions
    def_pattern = r'\\begin\{definisjon\}(.*?)\\end\{definisjon\}'
    for match in re.finditer(def_pattern, chunk, re.DOTALL):
        content = _clean_latex(match.group(1))
        para = doc.add_paragraph()
        run = para.add_run("📘 Definisjon: ")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 102, 204)
        para.add_run(content)
        para.style = 'Definition'
    
    # Remove processed definitions from chunk
    chunk = re.sub(def_pattern, '', chunk, flags=re.DOTALL)
    
    # Handle examples
    ex_pattern = r'\\begin\{eksempel\}(?:\[title=([^\]]+)\])?(.*?)\\end\{eksempel\}'
    for match in re.finditer(ex_pattern, chunk, re.DOTALL):
        title = match.group(1) or "Eksempel"
        content = _clean_latex(match.group(2))
        para = doc.add_paragraph()
        run = para.add_run(f"💡 {_clean_latex(title)}: ")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 153, 76)
        para.add_run(content)
        para.style = 'Example'
    
    # Remove processed examples
    chunk = re.sub(ex_pattern, '', chunk, flags=re.DOTALL)
    
    # Handle tasks
    task_pattern = r'\\begin\{taskbox\}\{([^}]+)\}(.*?)\\end\{taskbox\}'
    for match in re.finditer(task_pattern, chunk, re.DOTALL):
        title = _clean_latex(match.group(1))
        content = _clean_latex(match.group(2))
        para = doc.add_paragraph()
        run = para.add_run(f"✍️ {title}: ")
        run.bold = True
        para.add_run(content)
        para.style = 'Task'
    
    # Remove processed tasks
    chunk = re.sub(task_pattern, '', chunk, flags=re.DOTALL)
    
    # Handle merk/tips
    merk_pattern = r'\\begin\{merk\}(.*?)\\end\{merk\}'
    for match in re.finditer(merk_pattern, chunk, re.DOTALL):
        content = _clean_latex(match.group(1))
        para = doc.add_paragraph()
        run = para.add_run("💬 Merk: ")
        run.bold = True
        run.font.color.rgb = RGBColor(230, 126, 34)
        para.add_run(content)
    
    # Remove processed merk
    chunk = re.sub(merk_pattern, '', chunk, flags=re.DOTALL)
    
    # Handle løsning
    losning_pattern = r'\\begin\{losning\}(.*?)\\end\{losning\}'
    for match in re.finditer(losning_pattern, chunk, re.DOTALL):
        content = _clean_latex(match.group(1))
        para = doc.add_paragraph()
        run = para.add_run("🔑 Løsning: ")
        run.bold = True
        para.add_run(content)
    
    # Remove processed løsning
    chunk = re.sub(losning_pattern, '', chunk, flags=re.DOTALL)
    
    # Handle remaining content (paragraphs)
    remaining = chunk.strip()
    if remaining:
        # Split by double newlines or \par
        paragraphs = re.split(r'\n\n+|\\par\b', remaining)
        for para_text in paragraphs:
            para_text = _clean_latex(para_text).strip()
            if para_text and len(para_text) > 5:  # Skip very short fragments
                doc.add_paragraph(para_text)


def _clean_latex(text: str) -> str:
    """
    Clean LaTeX markup and convert to plain text.
    
    Args:
        text: LaTeX text.
    
    Returns:
        Cleaned text.
    """
    if not text:
        return ""
    
    # Remove comments
    text = re.sub(r'%.*$', '', text, flags=re.MULTILINE)
    
    # Handle text formatting
    text = re.sub(r'\\textbf\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\textit\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\emph\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\underline\{([^}]+)\}', r'\1', text)
    
    # Handle math - convert to readable format
    # Inline math
    def convert_math(match):
        math = match.group(1)
        # Simple conversions
        math = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1)/(\2)', math)
        math = re.sub(r'\\sqrt\{([^}]+)\}', r'√(\1)', math)
        math = re.sub(r'\\cdot', '·', math)
        math = re.sub(r'\\times', '×', math)
        math = re.sub(r'\\div', '÷', math)
        math = re.sub(r'\\pm', '±', math)
        math = re.sub(r'\\leq', '≤', math)
        math = re.sub(r'\\geq', '≥', math)
        math = re.sub(r'\\neq', '≠', math)
        math = re.sub(r'\\approx', '≈', math)
        math = re.sub(r'\\infty', '∞', math)
        math = re.sub(r'\\pi', 'π', math)
        math = re.sub(r'\\alpha', 'α', math)
        math = re.sub(r'\\beta', 'β', math)
        math = re.sub(r'\\theta', 'θ', math)
        math = re.sub(r'\\sum', 'Σ', math)
        math = re.sub(r'\\int', '∫', math)
        math = re.sub(r'\^(\d)', r'^\1', math)  # Keep superscripts
        math = re.sub(r'\^{([^}]+)}', r'^(\1)', math)
        math = re.sub(r'_(\d)', r'_\1', math)  # Keep subscripts
        math = re.sub(r'_{([^}]+)}', r'_(\1)', math)
        math = re.sub(r'\\left', '', math)
        math = re.sub(r'\\right', '', math)
        return math
    
    text = re.sub(r'\$([^$]+)\$', convert_math, text)
    
    # Handle display math
    text = re.sub(r'\\begin\{equation\*?\}', '', text)
    text = re.sub(r'\\end\{equation\*?\}', '', text)
    text = re.sub(r'\\begin\{align\*?\}', '', text)
    text = re.sub(r'\\end\{align\*?\}', '', text)
    text = re.sub(r'\\\[', '', text)
    text = re.sub(r'\\\]', '', text)
    
    # Handle lists
    text = re.sub(r'\\begin\{itemize\}', '', text)
    text = re.sub(r'\\end\{itemize\}', '', text)
    text = re.sub(r'\\begin\{enumerate\}(?:\[[^\]]*\])?', '', text)
    text = re.sub(r'\\end\{enumerate\}', '', text)
    text = re.sub(r'\\item\s*', '• ', text)
    
    # Handle tables (simplified)
    text = re.sub(r'\\begin\{table\}.*?\\end\{table\}', '[Tabell]', text, flags=re.DOTALL)
    text = re.sub(r'\\begin\{tabular\}.*?\\end\{tabular\}', '', text, flags=re.DOTALL)
    
    # Handle figures
    text = re.sub(r'\\begin\{figure\}.*?\\end\{figure\}', '[Figur]', text, flags=re.DOTALL)
    text = re.sub(r'\\begin\{tikzpicture\}.*?\\end\{tikzpicture\}', '[Graf/Figur]', text, flags=re.DOTALL)
    
    # Handle multicols
    text = re.sub(r'\\begin\{multicols\}\{\d+\}', '', text)
    text = re.sub(r'\\end\{multicols\}', '', text)
    
    # Remove other LaTeX commands
    text = re.sub(r'\\[a-zA-Z]+\*?(?:\[[^\]]*\])?(?:\{[^}]*\})*', '', text)
    
    # Clean up whitespace
    text = re.sub(r'\\\\', '\n', text)
    text = re.sub(r'&', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    text = text.strip()
    
    return text


def convert_latex_file_to_word(tex_path: str, output_dir: Optional[str] = None) -> Optional[str]:
    """
    Convert a LaTeX file to Word format.
    
    Args:
        tex_path: Path to the .tex file.
        output_dir: Output directory. Defaults to same as input.
    
    Returns:
        Path to the generated Word document.
    """
    tex_path = Path(tex_path)
    
    if not tex_path.exists():
        raise FileNotFoundError(f"LaTeX file not found: {tex_path}")
    
    with open(tex_path, "r", encoding="utf-8") as f:
        latex_content = f.read()
    
    if output_dir:
        output_path = Path(output_dir) / f"{tex_path.stem}.docx"
    else:
        output_path = tex_path.with_suffix('.docx')
    
    return latex_to_word(latex_content, str(output_path))
