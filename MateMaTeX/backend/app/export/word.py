"""
LaTeX → Word (DOCX) conversion.

Uses python-docx; inline math is rendered as native Word OMML when possible,
with Cambria Math fallback.
"""

from __future__ import annotations

import io
import re

import structlog

logger = structlog.get_logger()


_DOCX_SOLUTION_ENV = re.compile(r'\\begin\{losning\}.*?\\end\{losning\}', re.DOTALL)
_DOCX_SOLUTION_SECTION = re.compile(
    r'\\section\*?\{\s*L[øo]sning(?:sforslag)?\s*\}.*?(?=\\section|\\end\{document\}|\Z)',
    re.DOTALL | re.IGNORECASE,
)


def _strip_latex_commands(text: str, include_solutions: bool = True) -> str:
    """Convert LaTeX to readable plain text for Word export."""
    # Remove comments
    text = re.sub(r'%.*$', '', text, flags=re.MULTILINE)

    # Remove preamble
    doc_begin = text.find(r'\begin{document}')
    if doc_begin >= 0:
        text = text[doc_begin + len(r'\begin{document}'):]
    doc_end = text.find(r'\end{document}')
    if doc_end >= 0:
        text = text[:doc_end]

    # Optionally strip solutions for "elevkopi" mode.
    if not include_solutions:
        text = _DOCX_SOLUTION_ENV.sub('', text)
        text = _DOCX_SOLUTION_SECTION.sub('', text)

    # Remove environments (keep content)
    text = re.sub(r'\\begin\{[^}]*\}(?:\{[^}]*\})?', '', text)
    text = re.sub(r'\\end\{[^}]*\}', '', text)

    # Replace common commands
    replacements = {
        r'\maketitle': '',
        r'\newpage': '',
        r'\noindent': '',
        r'\vspace{': '',
        r'\hspace{': '',
        r'\\': '\n',
        r'\item': '• ',
        r'\textbf{': '',
        r'\textit{': '',
        r'\emph{': '',
        r'\sffamily': '',
        r'\bfseries': '',
    }
    for old, new in replacements.items():
        text = text.replace(old, new)

    # Extract section titles
    text = re.sub(r'\\section\*?\{([^}]*)\}', r'\n\n\1\n', text)
    text = re.sub(r'\\subsection\*?\{([^}]*)\}', r'\n\1\n', text)
    text = re.sub(r'\\title\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\author\{([^}]*)\}', r'\1', text)

    # Simplify math
    text = re.sub(r'\$\$([^$]*)\$\$', r' \1 ', text)
    text = re.sub(r'\$([^$]*)\$', r'\1', text)
    text = re.sub(r'\\\[([^]]*)\\\]', r' \1 ', text)
    text = re.sub(r'\\\(([^)]*)\\\)', r'\1', text)
    text = re.sub(r'\\frac\{([^}]*)\}\{([^}]*)\}', r'(\1)/(\2)', text)
    text = re.sub(r'\\sqrt\{([^}]*)\}', r'√(\1)', text)
    text = re.sub(r'\\cdot', r'·', text)
    text = re.sub(r'\\times', r'×', text)
    text = re.sub(r'\\pm', r'±', text)
    text = re.sub(r'\\leq', r'≤', text)
    text = re.sub(r'\\geq', r'≥', text)
    text = re.sub(r'\\neq', r'≠', text)
    text = re.sub(r'\\pi', r'π', text)
    text = re.sub(r'\\alpha', r'α', text)
    text = re.sub(r'\\beta', r'β', text)

    # Clean remaining commands
    text = re.sub(r'\\[a-zA-Z]+\*?\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\[a-zA-Z]+\*?', '', text)
    text = re.sub(r'[{}]', '', text)

    # Clean whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def _extract_body(latex_content: str, include_solutions: bool = True) -> str:
    """Document body with inline $...$ preserved for Word math runs."""
    text = re.sub(r'%.*$', '', latex_content, flags=re.MULTILINE)
    doc_begin = text.find(r"\begin{document}")
    if doc_begin >= 0:
        text = text[doc_begin + len(r"\begin{document}") :]
    doc_end = text.find(r"\end{document}")
    if doc_end >= 0:
        text = text[:doc_end]
    if not include_solutions:
        text = _DOCX_SOLUTION_ENV.sub("", text)
        text = _DOCX_SOLUTION_SECTION.sub("", text)
    text = re.sub(r"\\begin\{[^}]*\}(?:\{[^}]*\})?", "", text)
    text = re.sub(r"\\end\{[^}]*\}", "", text)
    text = re.sub(r"\\section\*?\{([^}]*)\}", r"\n\n\1\n", text)
    text = re.sub(r"\\subsection\*?\{([^}]*)\}", r"\n\1\n", text)
    return text.strip()


def _latex_inline_to_readable(expr: str) -> str:
    """Convert inline LaTeX to readable Unicode via MathML."""
    try:
        import xml.etree.ElementTree as ET

        from latex2mathml.converter import convert

        mathml = convert(expr.strip())
        plain = "".join(ET.fromstring(mathml).itertext())
        return plain.strip() or expr
    except Exception:
        return expr.replace("\\", "")


def _append_omml_inline(paragraph, latex_expr: str) -> bool:
    """Insert Office Math (OMML) for inline LaTeX; return False on failure."""
    try:
        from docx.oxml import parse_xml
        from latex2mathml.converter import convert as latex_to_mathml
        from mathml2omml import convert as mathml_to_omml

        omml = mathml_to_omml(latex_to_mathml(latex_expr.strip()))
        if not omml.strip().startswith("<m:oMath"):
            omml = (
                '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
                f"{omml}</m:oMath>"
            )
        paragraph._p.append(parse_xml(omml))
        return True
    except Exception:
        return False


def _add_word_paragraph(doc, line: str) -> None:
    """Paragraph with OMML (or Cambria Math fallback) for $...$ segments."""
    from docx.shared import Pt

    if line.startswith("• "):
        doc.add_paragraph(line[2:], style="List Bullet")
        return

    if "$" not in line:
        doc.add_paragraph(line)
        return

    p = doc.add_paragraph()
    parts = re.split(r"\$([^$]+)\$", line)
    for idx, part in enumerate(parts):
        if not part:
            continue
        if idx % 2 == 1:
            if not _append_omml_inline(p, part):
                run = p.add_run(_latex_inline_to_readable(part))
                run.italic = True
                run.font.name = "Cambria Math"
                run.font.size = Pt(11)
        else:
            p.add_run(part)


def latex_to_docx(
    latex_content: str,
    title: str = "Oppgaveark",
    include_solutions: bool = True,
) -> bytes:
    """
    Convert LaTeX content to a Word document.

    Returns the DOCX file as bytes.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Style defaults
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Generert av MateMaTeX AI")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # Spacer

    body = _extract_body(latex_content, include_solutions=include_solutions)
    body = re.sub(r"\$\$([^$]+)\$\$", r" $\1$ ", body)
    for section in re.split(r"\n\n+", body):
        section = section.strip()
        if not section:
            continue
        lines = section.split("\n")
        first_line = lines[0].strip()
        if len(lines) == 1 and len(first_line) < 80 and not first_line.startswith("•"):
            if any(
                word in first_line.lower()
                for word in ["oppgave", "eksempel", "definisjon", "løsning", "del 1", "del 2"]
            ):
                doc.add_heading(first_line, level=2)
                continue
        for line in lines:
            line = line.strip()
            if line:
                _add_word_paragraph(doc, line)

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
